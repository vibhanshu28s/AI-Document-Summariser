import fitz  # PyMuPDF
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def generate_final_matched_report(docx_input, pdf_source, output_name):
    doc = Document(docx_input)
    pdf_doc = fitz.open(pdf_source)

    # 1. Force Landscape Engineering Layout
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width

        # Set narrow 0.4" margins to maximize drawing area
        section.left_margin = section.right_margin = Inches(0.4)
        section.top_margin = section.bottom_margin = Inches(0.4)

    # 2. Correct Technical Text Reversals (OCR Artifacts)
    corrections = {
        "NROOITFCTUORNTSNOC": "NOT FOR CONSTRUCTION",
        "ROTCESAHPLA": "ALPHA SECTOR",
        "ROTCESATEB": "BETA SECTOR",
        "ROTCESAMMAG": "GAMMA SECTOR",
        "ETALPMET": "TEMPLATE",
        "POTFOOR": "ROOFTOP",
        "TSISNOC LL IW EPOCS EHT": "THE SCOPE WILL CONSIST",
        "POTFOOR_DC09_T&TA_NOKUE": "EUKON_AT&T_90CD_ROOFTOP"
    }

    for para in doc.paragraphs:
        for reversed_str, replacement in corrections.items():
            if reversed_str in para.text:
                para.text = para.text.replace(reversed_str, replacement)

    # 3. Restore High-Resolution Drawings
    # This replaces text placeholders with 300 DPI renderings from the PDF
    for para in doc.paragraphs:
        if "Drawing/Layout Snapshot" in para.text:
            try:
                page_info = para.text.split("Page ")[1].split()[0].strip()
                page_num = int(page_info) - 1

                # Render high-res snapshot (3x zoom)
                page = pdf_doc.load_page(page_num)
                pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
                img_path = f"img/final_render_p{page_num + 1}.png"
                pix.save(img_path)

                # Clear placeholder text and insert image
                para.text = ""
                run = para.add_run()
                run.add_picture(img_path, width=Inches(10.2))  # Standard landscape width
            except Exception as e:
                print(f"Failed to render drawing for {para.text}: {e}")

    # 4. Re-implement 2-Column Notes for Title Sheet
    if doc.sections:
        sectPr = doc.sections[0]._sectPr
        cols = sectPr.find(qn('w:cols'))
        if cols is None:
            cols = OxmlElement('w:cols')
            sectPr.insert(0, cols)
        cols.set(qn('w:num'), '2')  # Two-column layout
        cols.set(qn('w:sep'), '1')  # Divider line

    doc.save(output_name)
    pdf_doc.close()
    print(f"Successfully generated final matched doc: {output_name}")


# Execute
generate_final_matched_report(
    'output.docx',
    'Sample Insurance Certificate_0.pdf',
    'Report_MATCHED_FINAL.docx'
)