import streamlit as st
import docx
import os
import google.genai as genai
from dotenv import load_dotenv

# 1. Setup & Configuration
load_dotenv()
AI_API = os.getenv("GOV")
client = genai.Client(api_key=AI_API)

st.set_page_config(page_title="Document Table Extractor", page_icon="📄")
st.title("📄 AI Table Extractor")

# 2. File Uploader
uploaded_file = st.file_uploader("Upload your Word document", type=["docx"])

if uploaded_file is not None:
    # Process the document
    doc = docx.Document(uploaded_file)

    # Extracting Paragraphs
    full_text = [para.text for para in doc.paragraphs]

    # TIP: To help the AI, also extract raw table content
    # so it has the data it needs to "format" it for you.
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_data))

    store = '\n'.join(full_text)

    if st.button("Extract Tables with AI"):
        with st.spinner("Processing document..."):
            try:
                # Note: Using 'gemini-2.0-flash' as 2.5 is not yet a standard release
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=store,
                    config={
                        "system_instruction": "You are a Document Processing Agent. Identify and extract all tables. Format them clearly as tables."
                    }
                )

                # 3. Display the Response
                st.subheader("Extracted Tables")
                st.markdown(response.text)

            except Exception as e:
                st.error(f"An error occurred: {e}")