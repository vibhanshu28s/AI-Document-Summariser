import streamlit as st
from pathlib import Path
import os
import fitz  # PyMuPDF
import pdfplumber
import docx
from docx import Document
from docx.shared import Inches
from google import genai
from dotenv import load_dotenv

# Page Config
st.set_page_config(page_title="Document Summariser", page_icon="📄")

# Ensure directories exist
for folder in ['uploaded_files_dir', 'parser_file_output', 'summary_output']:
    Path(folder).mkdir(parents=True, exist_ok=True)

st.title("📄 Document Summariser")

load_dotenv()
AI_API = os.getenv("GEMINI_API_KEY")

# Initialize Session State to keep track of progress
if 'parsed_file_path' not in st.session_state:
    st.session_state.parsed_file_path = None
if 'extracted_text' not in st.session_state:
    st.session_state.extracted_text = None

# 1. FILE UPLOAD
uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")

if uploaded_file:
    save_path = Path('uploaded_files_dir', uploaded_file.name)
    with open(save_path, mode='wb') as f:
        f.write(uploaded_file.getbuffer())

    # BUTTON 1: PARSE PDF
    if st.button("Step 1: Parse Document"):
        with st.spinner("Converting PDF to Word (Text, Tables, & Drawings)..."):
            try:
                output_docx_path = Path("parser_file_output", f"{uploaded_file.name}_parsed.docx")
                pdf_images = fitz.open(str(save_path))
                word_doc = Document()
                full_text_list = []

                with pdfplumber.open(save_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        # Text Extraction
                        text = page.extract_text()
                        if text:
                            word_doc.add_paragraph(text)
                            full_text_list.append(text)

                        # Table Extraction
                        tables = page.extract_tables()
                        for table in tables:
                            if not table: continue
                            docx_table = word_doc.add_table(rows=len(table), cols=len(table[0]))
                            docx_table.style = 'Table Grid'
                            for r_idx, row in enumerate(table):
                                for c_idx, cell in enumerate(row):
                                    cell_val = str(cell) if cell else ""
                                    docx_table.cell(r_idx, c_idx).text = cell_val
                            full_text_list.append(f"[Table Data Page {i + 1}]")

                        # Image Snapshot
                        pix = pdf_images[i].get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                        img_path = f"temp_page_{i}.png"
                        pix.save(img_path)
                        word_doc.add_picture(img_path, width=Inches(5.5))
                        os.remove(img_path)

                word_doc.save(output_docx_path)
                pdf_images.close()

                # Save to session state
                st.session_state.parsed_file_path = str(output_docx_path)
                st.session_state.extracted_text = "\n".join(full_text_list)

                st.success("✅ Parsing Complete!")
            except Exception as e:
                st.error(f"Parsing failed: {e}")

    # 2. SHOW PARSED PREVIEW & DOWNLOAD
    if st.session_state.parsed_file_path:
        st.divider()
        st.subheader("📁 Parsed Document Ready")

        with open(st.session_state.parsed_file_path, "rb") as f:
            st.download_button(
                label="📥 Download Parsed Word File",
                data=f,
                file_name=f"{uploaded_file.name}_parsed.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with st.expander("Preview Extracted Text"):
            st.text(st.session_state.extracted_text[:2000] + "...")

        # BUTTON 2: SUMMARIZE (Only shows after parsing is done)
        st.info("Now, you can generate an AI Audit Summary based on the parsed content.")
        if st.button("Step 2: Generate AI Summary"):
            with st.spinner("Analyzing the document..."):
                try:
                    client = genai.Client(api_key=AI_API)

                    prompt = f"""
                    You are a document analysis engine.

                    Your task is NOT to summarize.
                    Your task is to extract and structure **ALL information present in the document** without omitting any detail.

                    Follow these strict rules:
                    - Do NOT summarize, compress, or simplify.
                    - Do NOT remove repetition if it exists in the document.
                    - Preserve original meaning, numbers, and wording as much as possible.
                    - If text is unclear, still extract it and mark it as "Unclear".
                    - If something is not present, write "Not stated".
                    - Maintain a professional, structured, audit-style format.

                    ---

                    ## 1. DOCUMENT CLASSIFICATION
                    - Document Type:
                    - Domain (Legal / Technical / Financial / Medical / General / etc.):
                    - Language:
                    - Format (Structured / Semi-structured / Unstructured / OCR-derived):

                    ---

                    ## 2. COMPLETE ENTITY EXTRACTION
                    Extract ALL entities mentioned:
                    - People (Names, roles, designations)
                    - Organizations / Companies
                    - Locations / Addresses / Sites
                    - IDs (Policy No, Invoice No, Contract ID, Case ID, Asset ID, etc.)
                    - Contact details (Emails, phone numbers, URLs)
                    - Dates (ALL dates with context)
                    - Any other identifiable named entity

                    ---

                    ## 3. FULL CONTENT STRUCTURE
                    Reconstruct the document into structured sections:

                    - Headings / Titles
                    - Subheadings
                    - Paragraphs
                    - Bullet points / Lists
                    - Tables (convert into readable structured format)
                    - Clauses / Sections / Articles

                    Preserve original hierarchy as much as possible.

                    ---

                    ## 4. COMPLETE DATA EXTRACTION
                    Extract ALL factual and numerical data:
                    - Monetary values
                    - Percentages
                    - Quantities
                    - Measurements (area, size, bandwidth, etc.)
                    - Time durations / deadlines
                    - Technical parameters
                    - Financial figures
                    - Any other measurable or factual data

                    ---

                    ## 5. ACTIONS, CONDITIONS & LOGIC
                    Extract ALL:
                    - Instructions
                    - Conditions (if/then logic)
                    - Rules
                    - Dependencies
                    - Obligations
                    - Constraints
                    - Exceptions
                    - Eligibility criteria

                    ---

                    ## 6. TABLE & STRUCTURED DATA RECONSTRUCTION
                    If tables exist:
                    - Recreate them in structured format (row-wise clearly)
                    - Preserve all columns and values
                    - Do not skip any cell

                    ---

                    ## 7. RAW TEXT PRESERVATION
                    Include a cleaned but complete version of the original text:
                    - Fix obvious OCR issues if needed
                    - Do NOT remove content
                    - Maintain order

                    ---

                    ## 8. ANOMALIES & UNCERTAINTIES
                    Highlight:
                    - Unclear text
                    - Missing values
                    - Conflicting information
                    - OCR errors
                    - Incomplete sentences

                    ---

                    ## OUTPUT FORMAT:
                    Return everything in clean, well-structured markdown with clear headings and bullet points.

                    ---

                    DOCUMENT CONTENT:
                    {st.session_state.extracted_text}
                    """

                    response = client.models.generate_content(
                        model="gemini-2.5-flash",
                        contents=prompt
                    )

                    summary_text = response.text

                    # Display Summary
                    st.divider()
                    st.markdown("### 🤖 AI Audit Summary")
                    st.markdown(summary_text)

                    # Save summary to file for download
                    summary_path = Path("summary_output", "Summary_Report.docx")
                    sum_doc = Document()
                    sum_doc.add_heading('AI Document Summary', 0)
                    sum_doc.add_paragraph(summary_text)
                    sum_doc.save(summary_path)

                    with open(summary_path, "rb") as sf:
                        st.download_button(
                            label="📥 Download Summary (.docx)",
                            data=sf,
                            file_name="Audit_Summary.docx"
                        )

                except Exception as e:
                    st.error(f"Summarization failed: {e}")