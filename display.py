import docx
from dotenv import load_dotenv
import os
import google.genai as genai
from docx import Document


load_dotenv()

AI_API = os.getenv("GOV")

doc = docx.Document("ds_ext/output_7.docx")

full_text = []
for para in doc.paragraphs:
    full_text.append(para.text)

store='\n'.join(full_text)

client = genai.Client(api_key=AI_API)


try:
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=store,
        config={
             "system_instruction": "You are a Document Processing Agent. Identify and extract all tables. Format them clearly as tables.",
            # "temperature": 0.1
        }
    )


    print(response.text)
    extracted_table = response.text

except Exception as e:
    print(f"An error occurred: {e}")

doc = Document()
doc.add_heading('Extracted Document Content', 0)

# We split the text by newlines to create proper paragraphs in Word
paragraphs = extracted_table.split('\n')
for p in paragraphs:
    if p.strip(): # Only add non-empty lines
        doc.add_paragraph(p)

doc.save("ds_tab/output_7_tables.docx")

print("✅ Saved to output.docx")