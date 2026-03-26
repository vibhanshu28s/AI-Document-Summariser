import docx
import os
from google import genai

import os
from dotenv import load_dotenv

load_dotenv()

AI_API = os.getenv("GEMINI_API_KEY")

# 1. Configuration

client = genai.Client(api_key=AI_API)


def extract_text_from_docx(file_path):
    if not os.path.exists(file_path):
        return None
    doc = docx.Document(file_path)
    full_content = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                full_content.append(" | ".join(row_data))
    return "\n".join(full_content)


def generate_summary(text):
    # Note: Ensure gemini-2.0-flash or gemini-1.5-flash is used if 2.5 isn't out yet!
    model_id = "gemini-2.5-flash"

    prompt = f"""
    You are a document analysis engine.

    Your task is NOT to summarize.
    Your task is to extract and structure **ALL information present in the document** without omitting any detail.

    Follow these strict rules:
    - Do NOT summarize, compress, or simplify.
    - Do NOT remove repetition if it exists in the document.
    - Preserve original meaning, numbers, and wording as much as possible.
    - If text is unclear, still extract it and mark it as "Unclear".
    - If something is not present, write "Not stated".
    - Maintain a professional, structured, audit-style format.

    ---

    ## 1. DOCUMENT CLASSIFICATION
    - Document Type:
    - Domain (Legal / Technical / Financial / Medical / General / etc.):
    - Language:
    - Format (Structured / Semi-structured / Unstructured / OCR-derived):

    ---

    ## 2. COMPLETE ENTITY EXTRACTION
    Extract ALL entities mentioned:
    - People (Names, roles, designations)
    - Organizations / Companies
    - Locations / Addresses / Sites
    - IDs (Policy No, Invoice No, Contract ID, Case ID, Asset ID, etc.)
    - Contact details (Emails, phone numbers, URLs)
    - Dates (ALL dates with context)
    - Any other identifiable named entity

    ---

    ## 3. FULL CONTENT STRUCTURE
    Reconstruct the document into structured sections:

    - Headings / Titles
    - Subheadings
    - Paragraphs
    - Bullet points / Lists
    - Tables (convert into readable structured format)
    - Clauses / Sections / Articles

    Preserve original hierarchy as much as possible.

    ---

    ## 4. COMPLETE DATA EXTRACTION
    Extract ALL factual and numerical data:
    - Monetary values
    - Percentages
    - Quantities
    - Measurements (area, size, bandwidth, etc.)
    - Time durations / deadlines
    - Technical parameters
    - Financial figures
    - Any other measurable or factual data

    ---

    ## 5. ACTIONS, CONDITIONS & LOGIC
    Extract ALL:
    - Instructions
    - Conditions (if/then logic)
    - Rules
    - Dependencies
    - Obligations
    - Constraints
    - Exceptions
    - Eligibility criteria

    ---

    ## 6. TABLE & STRUCTURED DATA RECONSTRUCTION
    If tables exist:
    - Recreate them in structured format (row-wise clearly)
    - Preserve all columns and values
    - Do not skip any cell

    ---

    ## 7. RAW TEXT PRESERVATION
    Include a cleaned but complete version of the original text:
    - Fix obvious OCR issues if needed
    - Do NOT remove content
    - Maintain order

    ---

    ## 8. ANOMALIES & UNCERTAINTIES
    Highlight:
    - Unclear text
    - Missing values
    - Conflicting information
    - OCR errors
    - Incomplete sentences

    ---

    ## OUTPUT FORMAT:
    Return everything in clean, well-structured markdown with clear headings and bullet points.

    ---

    DOCUMENT CONTENT:
    {text}
    """


    try:
        response = client.models.generate_content(
            model=model_id,
            contents=prompt
        )
        return response.text
    except Exception as e:
        return f"\nFinal Error: {e}"


def save_summary_to_docx(summary_text, output_path):
    """Saves the generated markdown summary into a formatted Word document."""
    doc = docx.Document()
    doc.add_heading('Document Summary', 0)

    # Simple parsing: break text into lines to handle basic formatting
    for line in summary_text.split('\n'):
        line = line.strip()
        if not line:
            continue

        if line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=1)
        elif line.startswith('-'):
            doc.add_paragraph(line.replace('-', '', 1).strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Summary saved successfully to: {output_path}")


# --- Execution ---
file_name = "input_extracted.docx"
output_name = "Summary_output.docx"
content = extract_text_from_docx(file_name)

if content:
    print(f"Summarizing {file_name}...")
    summary = generate_summary(content)

    if "Final Error" not in summary:
        print("Generation complete. Saving to file...")
        save_summary_to_docx(summary, output_name)
        print(summary)
    else:
        print(summary)
else:
    print("File not found.")