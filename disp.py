import streamlit as st
import json
import os
import io
from docx import Document
from google import genai
from google.genai import types
from dotenv import load_dotenv

# 1. Setup & Configuration
load_dotenv()
# Using one key for the combined app; fallback to "GOV" or "DISPLAY"
AI_API = os.getenv("GOV") or os.getenv("DISPLAY")
client = genai.Client(api_key=AI_API)

st.set_page_config(page_title="AI Document Digitizer", page_icon="📑", layout="wide")

## --- UI Header ---
st.title("📑 AI Document Digitizer & Table Extractor")
st.markdown("Upload a PDF to extract text and tables with high precision using Gemini 2.5-Flash.")

# 2. File Uploader
uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])

if uploaded_file is not None:
    # Read PDF bytes for Gemini
    pdf_bytes = uploaded_file.read()

    # Sidebar Options
    st.sidebar.header("Extraction Settings")
    maintain_layout = st.sidebar.checkbox("Maintain Layout/Tables", value=True)

    if st.button("🚀 Process & Extract Content"):
        with st.spinner("Gemini is digitizing the document..."):
            try:
                # 3. Generate Content using the OCR Logic
                prompt = "Scrape this entire PDF perfectly. Maintain the layout and tables." if maintain_layout else "Extract all text and tables from this document."

                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=[
                        prompt,
                        types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf"),
                    ],
                    config={
                        "system_instruction": "You are a High-Precision Document Digitization Expert. Identify and extract all text and tables. Format tables in Markdown.",
                    }
                )

                extracted_text = response.text

                # 4. Layout Columns for Results
                col1, col2 = st.columns(2)

                with col1:
                    st.subheader("📄 Extracted Content")
                    st.markdown(extracted_text)

                with col2:
                    st.subheader("💾 Export Options")

                    # --- Create Word Document in Memory ---
                    doc = Document()
                    doc.add_heading('Extracted Document Content', 0)
                    for p in extracted_text.split('\n'):
                        if p.strip():
                            doc.add_paragraph(p)

                    # Save doc to a buffer
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)

                    # --- Create JSON in Memory ---
                    json_data = json.dumps({
                        "source_file": uploaded_file.name,
                        "extraction_method": "gemini-2.5-flash",
                        "content": extracted_text
                    }, indent=4, ensure_ascii=False)

                    # Download Buttons
                    st.download_button(
                        label="Download as Word (.docx)",
                        data=doc_buffer,
                        file_name=f"{uploaded_file.name.rsplit('.', 1)[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    st.download_button(
                        label="Download as JSON",
                        data=json_data,
                        file_name="extraction_output.json",
                        mime="application/json"
                    )

            except Exception as e:
                st.error(f"An error occurred: {e}")
                st.info("Check if your API key is correct and has access to gemini-2.5-flash.")

else:
    st.info("Please upload a PDF file to begin.")