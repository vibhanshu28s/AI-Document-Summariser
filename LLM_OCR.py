import json
from google import genai
from google.genai import types
from docx import Document
from dotenv import load_dotenv
import os

load_dotenv()

AI_API = os.getenv("DISPLAY")


# Initialize Client
client = genai.Client(api_key=AI_API)

# 1. Load the PDF
pdf_path = "ds/input_7.pdf"
with open(pdf_path, "rb") as f:
    pdf_bytes = f.read()

# 2. Generate Content
# Prompting for Markdown helps maintain structure for the Word doc conversion
response = client.models.generate_content(
    model="gemini-2.5-flash",
    contents=[ """I want to convert pdf into document.
                "Scrape this entire PDF perfectly. Maintain the layout and tables.""",

            types.Part.from_bytes(data=pdf_bytes,mime_type="application/pdf"),
               ],
    config={
        "system_instruction": "You are a High-Precision Document Digitization Expert.",
        # "temperature": 0.7,           # Forces deterministic output for coordinates and IDs[cite: 12, 14, 46].
        # "top_p": 0.1,                # Ensures the model doesn't 'guess' characters in small technical fonts.
        # "max_output_tokens": 10000000,
        # "seed" : 42,

    }
)



extracted_text = response.text

# ---

### 3. Save as JSON
# We wrap the text in a dictionary to make it a valid JSON file
data_to_save = {
    "source_file": pdf_path,
    "extraction_method": "gemini-2.5-flash",
    "content": extracted_text
}

with open("output.json", "w", encoding="utf-8") as jf:
    json.dump(data_to_save, jf, indent=4, ensure_ascii=False)

print("✅ Saved to output.json")

# ---

### 4. Save as Word (.docx)
doc = Document()
doc.add_heading('Extracted Document Content', 0)

# We split the text by newlines to create proper paragraphs in Word
paragraphs = extracted_text.split('\n')
for p in paragraphs:
    if p.strip(): # Only add non-empty lines
        doc.add_paragraph(p)

doc.save("ds_ext/input_7.docx")

print("✅ Saved to output.docx")