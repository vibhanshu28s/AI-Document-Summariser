import json
from google import genai
from google.genai import types
from docx import Document
import os
from dotenv import load_dotenv


load_dotenv()

AI_API = os.getenv("GOV")

client = genai.Client(api_key=AI_API)


pdf_path = "ds/input_4.pdf"
with open(pdf_path, "rb") as f:
    pdf_bytes = f.read()

response = client.models.generate_content(
    model="gemini-2.5-flash",
    contents=[
        """Perform a high-fidelity OCR and layout extraction on this PDF. 

                                STRICT RULES:
                                1. STRUCTURE: Replicate the logical flow and hierarchy.
                                2. TABLES: Convert tables to Markdown.
                                3. NO CHATTER: Output ONLY document content.
                                4. NO LOOPING: Do not repeat segments of text. If you find yourself outputting the same line twice, move to the next section immediately.
                                5. PRECISION: Maintain original capitalization and punctuation.""",
        types.Part.from_bytes(
            data=pdf_bytes,
            mime_type="application/pdf"
        )
    ],
        config=types.GenerateContentConfig(
                        temperature=0.0,           # Increased slightly to avoid deterministic loops
                        # frequency_penalty=0.8,    # NEW: Penalizes the model for repeating tokens
                        # presence_penalty=0.3      # NEW: Encourages the model to move to new content
                    )
)

extracted_text = response.text

# ---

### 3. Save as JSON
# We wrap the text in a dictionary to make it a valid JSON file
data_to_save = {
    "source_file": pdf_path,
    "extraction_method": "gemini-2.5-flash",
    "content": extracted_text
}

with open("output.json", "w", encoding="utf-8") as jf:
    json.dump(data_to_save, jf, indent=4, ensure_ascii=False)

print("✅ Saved to output.json")

# ---

### 4. Save as Word (.docx)
doc = Document()
doc.add_heading('Extracted Document Content', 0)

# We split the text by newlines to create proper paragraphs in Word
paragraphs = extracted_text.split('\n')
for p in paragraphs:
    if p.strip(): # Only add non-empty lines
        doc.add_paragraph(p)

doc.save("ds_ext/output_input_4.docx")

print("✅ Saved to output.docx")