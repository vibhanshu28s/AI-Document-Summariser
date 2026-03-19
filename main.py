import streamlit as st
from pathlib import Path
import os
import fitz  # PyMuPDF
import pdfplumber
import docx
from docx import Document
from docx.shared import Inches
from google import genai
from dotenv import load_dotenv

# Page Config
st.set_page_config(page_title="Document Summariser", page_icon="📄")

# Ensure directories exist
for folder in ['uploaded_files_dir', 'parser_file_output', 'summary_output']:
    Path(folder).mkdir(parents=True, exist_ok=True)

st.title("📄 Document Summariser")

load_dotenv()
AI_API = os.getenv("GEMINI_API_KEY")

# Initialize Session State to keep track of progress
if 'parsed_file_path' not in st.session_state:
    st.session_state.parsed_file_path = None
if 'extracted_text' not in st.session_state:
    st.session_state.extracted_text = None

# 1. FILE UPLOAD
uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")

if uploaded_file:
    save_path = Path('uploaded_files_dir', uploaded_file.name)
    with open(save_path, mode='wb') as f:
        f.write(uploaded_file.getbuffer())

    # BUTTON 1: PARSE PDF
    if st.button("Step 1: Parse Document"):
        with st.spinner("Converting PDF to Word (Text, Tables, & Drawings)..."):
            try:
                output_docx_path = Path("parser_file_output", f"{uploaded_file.name}_parsed.docx")
                pdf_images = fitz.open(str(save_path))
                word_doc = Document()
                full_text_list = []

                with pdfplumber.open(save_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        # Text Extraction
                        text = page.extract_text()
                        if text:
                            word_doc.add_paragraph(text)
                            full_text_list.append(text)

                        # Table Extraction
                        tables = page.extract_tables()
                        for table in tables:
                            if not table: continue
                            docx_table = word_doc.add_table(rows=len(table), cols=len(table[0]))
                            docx_table.style = 'Table Grid'
                            for r_idx, row in enumerate(table):
                                for c_idx, cell in enumerate(row):
                                    cell_val = str(cell) if cell else ""
                                    docx_table.cell(r_idx, c_idx).text = cell_val
                            full_text_list.append(f"[Table Data Page {i + 1}]")

                        # Image Snapshot
                        pix = pdf_images[i].get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                        img_path = f"temp_page_{i}.png"
                        pix.save(img_path)
                        word_doc.add_picture(img_path, width=Inches(5.5))
                        os.remove(img_path)

                word_doc.save(output_docx_path)
                pdf_images.close()

                # Save to session state
                st.session_state.parsed_file_path = str(output_docx_path)
                st.session_state.extracted_text = "\n".join(full_text_list)

                st.success("✅ Parsing Complete!")
            except Exception as e:
                st.error(f"Parsing failed: {e}")

    # 2. SHOW PARSED PREVIEW & DOWNLOAD
    if st.session_state.parsed_file_path:
        st.divider()
        st.subheader("📁 Parsed Document Ready")

        with open(st.session_state.parsed_file_path, "rb") as f:
            st.download_button(
                label="📥 Download Parsed Word File",
                data=f,
                file_name=f"{uploaded_file.name}_parsed.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with st.expander("Preview Extracted Text"):
            st.text(st.session_state.extracted_text[:2000] + "...")

        # BUTTON 2: SUMMARIZE (Only shows after parsing is done)
        st.info("Now, you can generate an AI Audit Summary based on the parsed content.")
        if st.button("Step 2: Generate AI Summary"):
            with st.spinner("Analyzing the document..."):
                try:
                    client = genai.Client(api_key=AI_API)

                    prompt = f"""
                    You are a professional auditor. Summarize this document precisely.
                    Use headings for: Identity, Executive Summary (3 sentences), Critical Data, Obligations, Risks, and Final Assessment.

                    DOCUMENT CONTENT:
                    {st.session_state.extracted_text}
                    """

                    response = client.models.generate_content(
                        model="gemini-2.5-flash",
                        contents=prompt
                    )

                    summary_text = response.text

                    # Display Summary
                    st.divider()
                    st.markdown("### 🤖 AI Audit Summary")
                    st.markdown(summary_text)

                    # Save summary to file for download
                    summary_path = Path("summary_output", "Summary_Report.docx")
                    sum_doc = Document()
                    sum_doc.add_heading('AI Document Summary', 0)
                    sum_doc.add_paragraph(summary_text)
                    sum_doc.save(summary_path)

                    with open(summary_path, "rb") as sf:
                        st.download_button(
                            label="📥 Download Summary (.docx)",
                            data=sf,
                            file_name="Audit_Summary.docx"
                        )

                except Exception as e:
                    st.error(f"Summarization failed: {e}")